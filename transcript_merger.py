"""
transcript_merger.py
────────────────────────────────────────────────────────────────
Ericsson Transcript Merger — Core Processing Engine  v1.1
Zero Data Storage: all operations are purely in-memory.
No file is written to disk. No data is logged or transmitted.
────────────────────────────────────────────────────────────────
"""

import re
import io
from datetime import datetime, date
from typing import Optional


class TranscriptMerger:
    """
    Reads, parses, date-sorts, and merges multiple transcript files
    into a single coherent chronological document.

    Returns both plain-text (str) and Word (.docx bytes) outputs.

    Zero-storage guarantee:
        - All input is consumed from in-memory file objects.
        - All output is returned as Python str / bytes.
        - No disk I/O, no logging, no external calls.
    """

    DATE_PATTERNS = [
        (r'(\d{4})[\/\-](\d{1,2})[\/\-](\d{1,2})', '%Y%m%d'),
        (r'(\d{1,2})[\s\-\/]([A-Za-z]{3,9})[\s\-\/](\d{4})', 'dmy_word'),
        (r'([A-Za-z]{3,9})\s+(\d{1,2}),?\s+(\d{4})', 'mdy_word'),
        (r'(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{4})', 'ddmmyyyy'),
        (r'\b(\d{8})\b', 'compact'),
    ]

    MONTH_MAP = {
        'jan':1,'feb':2,'mar':3,'apr':4,'may':5,'jun':6,
        'jul':7,'aug':8,'sep':9,'oct':10,'nov':11,'dec':12,
        'january':1,'february':2,'march':3,'april':4,
        'june':6,'july':7,'august':8,'september':9,
        'october':10,'november':11,'december':12,
    }

    # ── Public API ──────────────────────────────────────────────────────────────

    def read_file(self, uploaded_file) -> str:
        name = uploaded_file.name.lower()
        raw_bytes = uploaded_file.getvalue()
        if name.endswith(".docx"):
            return self._read_docx_bytes(raw_bytes)
        for enc in ("utf-8","utf-8-sig","latin-1","cp1252"):
            try:
                return raw_bytes.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        return raw_bytes.decode("utf-8", errors="replace")

    def merge_transcripts(self, files_data: list[dict]) -> tuple[str, bytes, dict]:
        """
        Parameters
        ----------
        files_data : list of dicts with keys 'name' (str) and 'content' (str)

        Returns
        -------
        (merged_text: str, docx_bytes: bytes, stats: dict)
        """
        for record in files_data:
            record["date"] = self._extract_date(record["name"], record["content"])

        def sort_key(r):
            d = r["date"]
            return (0, d) if d else (1, date.max)

        sorted_records = sorted(files_data, key=sort_key)

        # ── Build plain-text ────────────────────────────────────────────────────
        sections = []
        for idx, record in enumerate(sorted_records, start=1):
            date_str = (
                record["date"].strftime("%d %B %Y")
                if record["date"] else "Date Not Detected"
            )
            divider = "═" * 72
            header = (
                f"\n{divider}\n"
                f"  TRANSCRIPT {idx} OF {len(sorted_records)}\n"
                f"  Source File : {record['name']}\n"
                f"  Meeting Date: {date_str}\n"
                f"{divider}\n\n"
            )
            sections.append(header + record["content"].strip())

        now_str = datetime.now().strftime("%d %B %Y, %H:%M")
        first_date = sorted_records[0]["date"].strftime("%d %b %Y") if sorted_records[0]["date"] else "Unknown"
        last_date  = sorted_records[-1]["date"].strftime("%d %b %Y") if sorted_records[-1]["date"] else "Unknown"

        doc_header = (
            "╔══════════════════════════════════════════════════════════════════════╗\n"
            "║           ERICSSON — MERGED TRANSCRIPT DOCUMENT                     ║\n"
            f"║           Generated: {now_str:<44}║\n"
            "║           Zero Data Storage | In-Memory Processing Only             ║\n"
            "╚══════════════════════════════════════════════════════════════════════╝\n\n"
            "SUMMARY\n"
            "───────\n"
            f"  Total transcripts merged : {len(sorted_records)}\n"
            f"  Date range               : {first_date} → {last_date}\n"
            "  Sort order               : Chronological (oldest first)\n\n"
        )

        merged_text = doc_header + "\n\n".join(sections)

        # ── Build DOCX ──────────────────────────────────────────────────────────
        docx_bytes = self._build_docx(sorted_records, now_str, first_date, last_date)

        # ── Stats ───────────────────────────────────────────────────────────────
        dated_count = sum(1 for r in sorted_records if r["date"])
        stats = {
            "total_files": len(sorted_records),
            "total_lines": sum(len(r["content"].splitlines()) for r in sorted_records),
            "total_words": sum(len(r["content"].split()) for r in sorted_records),
            "date_sorted": f"{dated_count}/{len(sorted_records)}",
        }

        return merged_text, docx_bytes, stats

    # ── DOCX Builder ────────────────────────────────────────────────────────────

    def _build_docx(
        self,
        sorted_records: list[dict],
        now_str: str,
        first_date: str,
        last_date: str,
    ) -> bytes:
        """Build a formatted Word document entirely in memory and return bytes."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import copy

            doc = Document()

            # ── Page margins ──
            for section in doc.sections:
                section.top_margin    = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin   = Inches(1.2)
                section.right_margin  = Inches(1.2)

            # ── Helper: set paragraph shading ──
            def set_para_shading(para, fill_hex: str):
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fill_hex)
                pPr.append(shd)

            # ── Helper: horizontal rule ──
            def add_hr(doc, color_hex="003366"):
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), color_hex)
                pBdr.append(bottom)
                pPr.append(pBdr)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                return p

            # ── COVER / HEADER ──────────────────────────────────────────────────
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after  = Pt(4)
            set_para_shading(title_para, "003366")
            run = title_para.add_run("  ERICSSON — MERGED TRANSCRIPT DOCUMENT")
            run.bold      = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = "Arial"

            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            subtitle_para.paragraph_format.space_before = Pt(0)
            subtitle_para.paragraph_format.space_after  = Pt(0)
            set_para_shading(subtitle_para, "003366")
            sub_run = subtitle_para.add_run(f"  Generated: {now_str}   |   Zero Data Storage — In-Memory Processing Only")
            sub_run.font.size  = Pt(9)
            sub_run.font.color.rgb = RGBColor(0xB0, 0xD0, 0xFF)
            sub_run.font.name  = "Arial"

            doc.add_paragraph()  # spacer

            # ── SUMMARY BOX ────────────────────────────────────────────────────
            sum_head = doc.add_paragraph()
            sum_head.paragraph_format.space_after = Pt(2)
            r = sum_head.add_run("SUMMARY")
            r.bold = True; r.font.size = Pt(10); r.font.name = "Arial"
            r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

            add_hr(doc)

            for line in [
                f"Total transcripts merged : {len(sorted_records)}",
                f"Date range               : {first_date}  →  {last_date}",
                "Sort order               : Chronological (oldest first)",
            ]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)
                r = p.add_run(line)
                r.font.size = Pt(9.5); r.font.name = "Courier New"
                r.font.color.rgb = RGBColor(0x2D, 0x3F, 0x55)

            doc.add_paragraph()

            # ── TRANSCRIPT SECTIONS ─────────────────────────────────────────────
            for idx, record in enumerate(sorted_records, start=1):
                date_str = (
                    record["date"].strftime("%d %B %Y")
                    if record["date"] else "Date Not Detected"
                )

                # Section header bar
                hdr_para = doc.add_paragraph()
                hdr_para.paragraph_format.space_before = Pt(12)
                hdr_para.paragraph_format.space_after  = Pt(0)
                set_para_shading(hdr_para, "0056A8")
                h_run = hdr_para.add_run(
                    f"  TRANSCRIPT {idx} OF {len(sorted_records)}   |   {date_str}"
                )
                h_run.bold = True; h_run.font.size = Pt(10.5); h_run.font.name = "Arial"
                h_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                # Source file sub-bar
                src_para = doc.add_paragraph()
                src_para.paragraph_format.space_before = Pt(0)
                src_para.paragraph_format.space_after  = Pt(6)
                set_para_shading(src_para, "E8F4FD")
                s_run = src_para.add_run(f"  Source: {record['name']}")
                s_run.font.size = Pt(8.5); s_run.font.name = "Arial"
                s_run.font.color.rgb = RGBColor(0x00, 0x56, 0xA8)

                # Body text — paragraph by paragraph
                for line in record["content"].strip().splitlines():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(1)
                    r = p.add_run(line)
                    r.font.size = Pt(10); r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

                # Separator
                if idx < len(sorted_records):
                    add_hr(doc, "BAC8D8")
                    doc.add_paragraph()

            # ── FOOTER ─────────────────────────────────────────────────────────
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_shading(footer_para, "0A1F3A")
            f_run = footer_para.add_run(
                "Ericsson Internal Tool   |   Transcript Merger v1.1   |   Zero Data Storage"
            )
            f_run.font.size = Pt(8); f_run.font.name = "Arial"
            f_run.font.color.rgb = RGBColor(0xB0, 0xC4, 0xD8)

            # ── Serialize to bytes ──────────────────────────────────────────────
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        except ImportError:
            # python-docx not installed — return a minimal valid stub
            return self._docx_not_available_stub()
        except Exception as e:
            return self._docx_not_available_stub(str(e))

    def _docx_not_available_stub(self, error: str = "") -> bytes:
        """Return a tiny text file disguised as docx error note."""
        msg = (
            "python-docx is not installed or an error occurred.\n"
            "Run: pip install python-docx\n"
        )
        if error:
            msg += f"\nError: {error}\n"
        # Return as plain bytes — Streamlit will still offer download
        return msg.encode("utf-8")

    # ── Private helpers ─────────────────────────────────────────────────────────

    def _read_docx_bytes(self, raw_bytes: bytes) -> str:
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            return "[⚠️ python-docx not installed. Run: pip install python-docx]"
        except Exception as e:
            return f"[⚠️ Could not parse .docx file: {e}]"

    def _extract_date(self, filename: str, content: str) -> Optional[date]:
        search_text = filename + "\n" + "\n".join(content.splitlines()[:50])
        for pattern, fmt in self.DATE_PATTERNS:
            match = re.search(pattern, search_text, re.IGNORECASE)
            if match:
                parsed = self._parse_match(match, fmt)
                if parsed:
                    return parsed
        return None

    def _parse_match(self, match: re.Match, fmt: str) -> Optional[date]:
        try:
            g = match.groups()
            if fmt == '%Y%m%d':
                return date(int(g[0]), int(g[1]), int(g[2]))
            if fmt == 'compact':
                s = g[0]
                return date(int(s[:4]), int(s[4:6]), int(s[6:8]))
            if fmt == 'dmy_word':
                d, m_str, y = int(g[0]), g[1].lower()[:3], int(g[2])
                m = self.MONTH_MAP.get(m_str)
                return date(y, m, d) if m else None
            if fmt == 'mdy_word':
                m_str, d, y = g[0].lower()[:3], int(g[1]), int(g[2])
                m = self.MONTH_MAP.get(m_str)
                return date(y, m, d) if m else None
            if fmt == 'ddmmyyyy':
                d, m, y = int(g[0]), int(g[1]), int(g[2])
                return date(y, m, d)
        except (ValueError, TypeError):
            return None
        return None
